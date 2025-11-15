import os
import json
from datetime import datetime
import flet as ft
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import urllib.request
import tkinter as tk
from tkinter import filedialog


class MireaReportGenerator:
    """Главный класс приложения - Генератор отчётов РТУ МИРЭА"""

    def __init__(self, page: ft.Page):
        self.page = page
        self.page.title = "MIREA Report Generator - Генератор отчётов РТУ МИРЭА"
        self.page.window.width = 900
        self.page.window.height = 800
        self.page.padding = 20
        self.page.scroll = ft.ScrollMode.ADAPTIVE

        self.page.locale_configuration = ft.LocaleConfiguration(
            supported_locales=[ft.Locale("ru", "RU")],
            current_locale=ft.Locale("ru", "RU")
        )

        self.selected_directory = None
        self.selected_save_directory = None
        self.found_files = []
        self.config_file = "config.json"
        self.selected_date = datetime.now()
        self.template_url = "https://raw.githubusercontent.com/Vennilay/MireaReportGenerator/main/template.docx"
        self.avatar_url = "https://avatars.githubusercontent.com/Vennilay"
        self.repo_url = "https://github.com/Vennilay/MireaReportGenerator"

        self.group_field = None
        self.student_field = None
        self.teacher_field = None
        self.work_number_field = None
        self.template_path_field = None
        self.date_picker = None
        self.date_display = None
        self.directory_text = None
        self.save_directory_text = None
        self.template_path_display = None
        self.save_nearby_checkbox = None
        self.files_count_text = None
        self.show_files_btn = None
        self.generate_btn = None
        self.select_save_dir_btn = None

        self.config = self.load_config()
        self.create_ui()

    def load_config(self):
        """Загружает сохранённые настройки из config.json"""
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            self.show_snackbar(f"Ошибка загрузки конфига: {str(e)}", ft.Colors.ORANGE)

        return {
            "group": "",
            "student_name": "",
            "teacher_name": "",
            "work_number": "",
            "last_directory": "",
            "template_path": "template.docx",
            "save_directory": "",
            "save_nearby": True
        }

    def save_config(self):
        """Сохраняет текущие настройки в config.json"""
        try:
            config = {
                "group": self.group_field.value,
                "student_name": self.student_field.value,
                "teacher_name": self.teacher_field.value,
                "work_number": self.work_number_field.value,
                "last_directory": self.selected_directory or "",
                "template_path": self.template_path_field.value,
                "save_directory": self.selected_save_directory or "",
                "save_nearby": self.save_nearby_checkbox.value
            }
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config, indent=4, fp=f, ensure_ascii=False)
        except Exception as e:
            self.show_snackbar(f"Ошибка сохранения конфига: {str(e)}", ft.Colors.ORANGE)

    def validate_form(self):
        """Проверяет заполнение всех обязательных полей и обновляет состояние кнопки"""
        if not self.generate_btn:
            return

        is_valid = (
                bool(self.group_field.value and self.group_field.value.strip()) and
                bool(self.student_field.value and self.student_field.value.strip()) and
                bool(self.teacher_field.value and self.teacher_field.value.strip()) and
                bool(self.work_number_field.value and self.work_number_field.value.strip()) and
                bool(self.found_files)
        )

        self.generate_btn.disabled = not is_valid

        if is_valid:
            self.generate_btn.bgcolor = ft.Colors.GREEN_700
            self.generate_btn.color = ft.Colors.WHITE
            self.generate_btn.opacity = 1.0
        else:
            self.generate_btn.bgcolor = ft.Colors.GREY_400
            self.generate_btn.color = ft.Colors.GREY_700
            self.generate_btn.opacity = 0.6

        self.generate_btn.update()

    def on_save_nearby_changed(self, _e):
        """Обработчик изменения чекбокса 'Сохранить рядом'"""
        if self.save_nearby_checkbox.value:
            self.select_save_dir_btn.disabled = True
            self.save_directory_text.value = "Файл будет сохранён рядом с программой"
            self.save_directory_text.color = ft.Colors.GREY_600
        else:
            self.select_save_dir_btn.disabled = False
            if self.selected_save_directory:
                self.save_directory_text.value = f"Папка сохранения: {self.selected_save_directory}"
                self.save_directory_text.color = ft.Colors.GREEN_700
            else:
                self.save_directory_text.value = "Выберите папку для сохранения"
                self.save_directory_text.color = ft.Colors.ORANGE_700

        self.page.update()

    def select_directory_tkinter(self, _e):
        """Выбор директории с кодом через tkinter"""
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        directory = filedialog.askdirectory(
            title="Выберите папку с файлами кода",
            initialdir=self.config.get("last_directory", "")
        )

        root.destroy()

        if directory:
            self.selected_directory = directory
            self.directory_text.value = f"Выбрана: {self.selected_directory}"
            self.directory_text.color = ft.Colors.GREEN_700
            self.find_code_files()
            self.page.update()
            self.show_snackbar(
                f"✅ Найдено файлов: {len(self.found_files)}",
                ft.Colors.GREEN_700
            )

    def select_template_tkinter(self, _e):
        """Выбор файла шаблона через tkinter"""
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        template_path = filedialog.askopenfilename(
            title="Выберите файл шаблона DOCX",
            filetypes=[("Word Documents", "*.docx"), ("All files", "*.*")],
            initialdir=os.path.dirname(self.template_path_field.value) if self.template_path_field.value else ""
        )

        root.destroy()

        if template_path:
            self.template_path_field.value = template_path
            self.template_path_display.value = f"Текущий шаблон: {os.path.basename(template_path)}"
            self.template_path_display.color = ft.Colors.GREEN_700
            self.page.update()
            self.show_snackbar(
                f"✅ Выбран шаблон: {os.path.basename(template_path)}",
                ft.Colors.GREEN_700
            )

    def select_save_directory_tkinter(self, _e):
        """Выбор папки для сохранения через tkinter"""
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        save_directory = filedialog.askdirectory(
            title="Выберите папку для сохранения документа",
            initialdir=self.config.get("save_directory", "")
        )

        root.destroy()

        if save_directory:
            self.selected_save_directory = save_directory
            self.save_directory_text.value = f"Папка сохранения: {self.selected_save_directory}"
            self.save_directory_text.color = ft.Colors.GREEN_700
            self.page.update()
            self.show_snackbar(
                f"✅ Выбрана папка: {self.selected_save_directory}",
                ft.Colors.GREEN_700
            )

    def download_template(self, _e):
        """Скачивает шаблон с GitHub"""
        try:
            self.show_snackbar("⏳ Скачивание шаблона с GitHub...", ft.Colors.BLUE_700)

            output_path = "template.docx"
            urllib.request.urlretrieve(self.template_url, output_path)

            self.template_path_field.value = output_path
            self.template_path_display.value = f"Текущий шаблон: {output_path}"
            self.template_path_display.color = ft.Colors.GREEN_700

            self.page.update()
            self.show_snackbar(
                "✅ Шаблон успешно скачан с GitHub!",
                ft.Colors.GREEN_700
            )
            self.show_dialog(
                "Успех! 🎉",
                f"Шаблон успешно скачан!\n\n"
                f"Файл сохранён как: {output_path}\n\n"
                f"Теперь вы можете использовать его для создания документов."
            )

        except Exception as e:
            error_message = f"Не удалось скачать шаблон:\n\n{str(e)}\n\nПроверьте подключение к интернету."
            self.show_dialog("Ошибка", error_message)
            self.show_snackbar(f"❌ Ошибка скачивания: {str(e)}", ft.Colors.RED_700)

    def show_snackbar(self, message: str, color: str = ft.Colors.BLUE_700):
        """Показывает уведомление в верхней части экрана (не перекрывает контент)"""
        snackbar = ft.SnackBar(
            content=ft.Text(message, color=ft.Colors.WHITE, size=14),
            bgcolor=color,
            duration=3000,
            behavior=ft.SnackBarBehavior.FLOATING,
            margin=ft.margin.only(top=10, left=20, right=20),
        )

        self.page.overlay.insert(0, snackbar)
        snackbar.open = True
        self.page.update()

    def show_dialog(self, title: str, message: str):
        """Показывает диалоговое окно с информацией или ошибкой"""

        def close_dialog(_e):
            dialog.open = False
            self.page.update()

        dialog = ft.AlertDialog(
            modal=True,
            title=ft.Text(title, weight=ft.FontWeight.BOLD),
            content=ft.Text(message),
            actions=[
                ft.TextButton("ОК", on_click=close_dialog)
            ],
            actions_alignment=ft.MainAxisAlignment.END
        )

        self.page.overlay.append(dialog)
        dialog.open = True
        self.page.update()

    def show_files_dialog(self, _e):
        """Показывает диалоговое окно со списком всех файлов"""

        def close_dialog(_e):
            dialog.open = False
            self.page.update()

        files_list_view = ft.ListView(
            spacing=5,
            padding=10,
            height=400,
            width=600
        )

        if self.found_files:
            files_list_view.controls.append(
                ft.Container(
                    content=ft.Text(
                        f"📁 Найдено файлов: {len(self.found_files)}",
                        weight=ft.FontWeight.BOLD,
                        color=ft.Colors.WHITE,
                        size=16
                    ),
                    bgcolor=ft.Colors.BLUE_700,
                    padding=12,
                    border_radius=8
                )
            )

            for idx, file_path in enumerate(self.found_files, 1):
                files_list_view.controls.append(
                    ft.Container(
                        content=ft.Row([
                            ft.Text(
                                f"{idx}.",
                                size=14,
                                color=ft.Colors.BLUE_400,
                                weight=ft.FontWeight.BOLD,
                                width=40
                            ),
                            ft.Icon(ft.Icons.CODE, size=20, color=ft.Colors.BLUE_300),
                            ft.Column([
                                ft.Text(
                                    os.path.basename(file_path),
                                    size=14,
                                    color=ft.Colors.WHITE,
                                    weight=ft.FontWeight.W_500
                                ),
                                ft.Text(
                                    os.path.dirname(file_path),
                                    size=11,
                                    color=ft.Colors.GREY_400,
                                    italic=True
                                )
                            ], spacing=2, tight=True)
                        ], spacing=10),
                        padding=12,
                        border_radius=6,
                        bgcolor=ft.Colors.GREY_800,
                        border=ft.border.all(1, ft.Colors.GREY_700)
                    )
                )
        else:
            files_list_view.controls.append(
                ft.Container(
                    content=ft.Row([
                        ft.Icon(ft.Icons.WARNING_AMBER, color=ft.Colors.ORANGE_400, size=28),
                        ft.Text(
                            "Файлы с кодом не найдены",
                            color=ft.Colors.ORANGE_300,
                            weight=ft.FontWeight.W_500,
                            size=14
                        )
                    ], spacing=10),
                    padding=20,
                    bgcolor=ft.Colors.GREY_800,
                    border_radius=8
                )
            )

        dialog = ft.AlertDialog(
            modal=True,
            title=ft.Text(
                "Список найденных файлов",
                weight=ft.FontWeight.BOLD,
                size=20
            ),
            content=ft.Container(
                content=files_list_view,
                bgcolor=ft.Colors.GREY_900,
                border_radius=8,
                padding=10
            ),
            actions=[
                ft.TextButton("Закрыть", on_click=close_dialog)
            ],
            actions_alignment=ft.MainAxisAlignment.END
        )

        self.page.overlay.append(dialog)
        dialog.open = True
        self.page.update()

    def show_about_dialog(self, _e):
        """Показывает информацию о создателе с аватаркой"""

        def close_dialog(_e):
            dialog.open = False
            self.page.update()

        def open_github(_e):
            self.page.launch_url("https://github.com/Vennilay")

        def open_repo(_e):
            self.page.launch_url(self.repo_url)

        dialog = ft.AlertDialog(
            modal=True,
            title=ft.Text("О создателе 👨‍💻", weight=ft.FontWeight.BOLD, size=20),
            content=ft.Column([
                ft.Container(
                    content=ft.CircleAvatar(
                        foreground_image_src=self.avatar_url,
                        radius=50,
                        bgcolor=ft.Colors.BLUE_700
                    ),
                    alignment=ft.alignment.center,
                    padding=ft.padding.only(bottom=15)
                ),

                ft.Text(
                    "MIREA Report Generator",
                    size=16,
                    weight=ft.FontWeight.BOLD,
                    color=ft.Colors.BLUE_700,
                    text_align=ft.TextAlign.CENTER
                ),
                ft.Divider(height=10),
                ft.Text(
                    "Генератор отчётов для студентов РТУ МИРЭА",
                    size=14,
                    color=ft.Colors.GREY_700,
                    text_align=ft.TextAlign.CENTER
                ),
                ft.Container(height=10),
                ft.Text(
                    "Разработчик: Vennilay",
                    size=14,
                    weight=ft.FontWeight.W_500,
                    text_align=ft.TextAlign.CENTER
                ),
                ft.Container(height=10),

                ft.Row([
                    ft.ElevatedButton(
                        "GitHub Profile",
                        icon=ft.Icons.PERSON,
                        on_click=open_github,
                        style=ft.ButtonStyle(
                            bgcolor=ft.Colors.GREY_800,
                            color=ft.Colors.WHITE
                        )
                    ),
                    ft.ElevatedButton(
                        "Repository",
                        icon=ft.Icons.CODE,
                        on_click=open_repo,
                        style=ft.ButtonStyle(
                            bgcolor=ft.Colors.BLUE_700,
                            color=ft.Colors.WHITE
                        )
                    ),
                ], alignment=ft.MainAxisAlignment.CENTER, spacing=10),

                ft.Container(height=10),
                ft.Text(
                    "© 2025 Vennilay",
                    size=12,
                    color=ft.Colors.GREY_500,
                    italic=True,
                    text_align=ft.TextAlign.CENTER
                )
            ], tight=True, spacing=5, horizontal_alignment=ft.CrossAxisAlignment.CENTER),
            actions=[
                ft.TextButton("Закрыть", on_click=close_dialog)
            ],
            actions_alignment=ft.MainAxisAlignment.END
        )

        self.page.overlay.append(dialog)
        dialog.open = True
        self.page.update()

    def create_ui(self):
        """Создаёт пользовательский интерфейс"""

        header_row = ft.Row([
            ft.Text(
                "MIREA Report Generator",
                size=26,
                weight=ft.FontWeight.BOLD,
                color=ft.Colors.BLUE_700
            ),
            ft.IconButton(
                icon=ft.Icons.INFO_OUTLINED,
                tooltip="О создателе",
                on_click=self.show_about_dialog,
                icon_color=ft.Colors.BLUE_600,
                icon_size=28
            )
        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN)

        subtitle = ft.Text(
            "Генератор отчётов для РТУ МИРЭА",
            size=14,
            color=ft.Colors.GREY_600,
            italic=True
        )

        self.group_field = ft.TextField(
            label="Группа (например: ИКБО-47-52)",
            value=self.config.get("group", ""),
            width=400,
            autofocus=True,
            border_color=ft.Colors.BLUE_400,
            prefix_icon=ft.Icons.GROUP,
            on_change=lambda _: self.validate_form()
        )

        self.student_field = ft.TextField(
            label="ФИО студента (например: Иванов И.И.)",
            value=self.config.get("student_name", ""),
            width=400,
            border_color=ft.Colors.BLUE_400,
            prefix_icon=ft.Icons.PERSON,
            on_change=lambda _: self.validate_form()
        )

        self.teacher_field = ft.TextField(
            label="ФИО преподавателя (например: Кодабашян Л.С.)",
            value=self.config.get("teacher_name", ""),
            width=400,
            border_color=ft.Colors.BLUE_400,
            prefix_icon=ft.Icons.SCHOOL,
            on_change=lambda _: self.validate_form()
        )

        self.work_number_field = ft.TextField(
            label="Номер работы",
            value=self.config.get("work_number", ""),
            width=200,
            keyboard_type=ft.KeyboardType.NUMBER,
            border_color=ft.Colors.BLUE_400,
            prefix_icon=ft.Icons.NUMBERS,
            on_change=lambda _: self.validate_form()
        )

        self.template_path_field = ft.TextField(
            label="Путь к файлу шаблона (например: template.docx)",
            value=self.config.get("template_path", "template.docx"),
            width=400,
            border_color=ft.Colors.BLUE_400,
            prefix_icon=ft.Icons.DESCRIPTION,
            hint_text="Укажите путь или имя файла шаблона"
        )

        # TKINTER КНОПКА вместо FilePicker
        select_template_btn = ft.ElevatedButton(
            "Выбрать файл",
            icon=ft.Icons.FILE_OPEN,
            on_click=self.select_template_tkinter,  # Теперь через tkinter
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.PURPLE_600,
                color=ft.Colors.WHITE
            )
        )

        download_template_btn = ft.ElevatedButton(
            "Скачать с GitHub",
            icon=ft.Icons.DOWNLOAD,
            on_click=self.download_template,
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.GREEN_600,
                color=ft.Colors.WHITE
            ),
            tooltip="Скачать шаблон template.docx с GitHub"
        )

        self.template_path_display = ft.Text(
            value=f"Текущий шаблон: {self.config.get('template_path', 'template.docx')}",
            color=ft.Colors.GREY_700,
            size=12
        )

        self.date_picker = ft.DatePicker(
            first_date=datetime(2020, 1, 1),
            last_date=datetime(2030, 12, 31),
            on_change=self.on_date_changed,
            on_dismiss=self.on_date_dismissed,
            help_text="Выберите дату",
            cancel_text="Отмена",
            confirm_text="ОК",
            error_format_text="Неверный формат",
            error_invalid_text="Вне диапазона",
            field_label_text="Введите дату",
            field_hint_text="дд.мм.гггг"
        )
        self.page.overlay.append(self.date_picker)

        self.date_display = ft.Text(
            value=self.format_date(self.selected_date),
            size=16,
            color=ft.Colors.GREEN_700,
            weight=ft.FontWeight.BOLD
        )

        date_picker_btn = ft.ElevatedButton(
            "Выбрать дату",
            icon=ft.Icons.CALENDAR_MONTH,
            on_click=self.open_date_picker,
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.BLUE_600,
                color=ft.Colors.WHITE
            )
        )

        self.directory_text = ft.Text(
            value="Директория не выбрана",
            color=ft.Colors.GREY_700
        )

        # TKINTER КНОПКА вместо FilePicker
        select_dir_btn = ft.ElevatedButton(
            "Выбрать директорию с кодом",
            icon=ft.Icons.FOLDER_OPEN,
            on_click=self.select_directory_tkinter,  # Теперь через tkinter
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.BLUE_600,
                color=ft.Colors.WHITE
            )
        )

        self.files_count_text = ft.Text(
            value="Файлы не найдены",
            color=ft.Colors.GREY_600,
            size=14
        )

        self.show_files_btn = ft.ElevatedButton(
            "Показать список файлов",
            icon=ft.Icons.LIST,
            on_click=self.show_files_dialog,
            visible=False,
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.INDIGO_600,
                color=ft.Colors.WHITE
            )
        )

        self.save_nearby_checkbox = ft.Checkbox(
            label="Сохранить рядом с программой",
            value=self.config.get("save_nearby", True),
            on_change=self.on_save_nearby_changed,
            fill_color=ft.Colors.BLUE_600
        )

        # TKINTER КНОПКА вместо FilePicker
        self.select_save_dir_btn = ft.ElevatedButton(
            "Выбрать папку для сохранения",
            icon=ft.Icons.FOLDER_SPECIAL,
            on_click=self.select_save_directory_tkinter,  # Теперь через tkinter
            disabled=self.config.get("save_nearby", True),
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.TEAL_600,
                color=ft.Colors.WHITE
            )
        )

        if self.config.get("save_nearby", True):
            initial_save_text = "Файл будет сохранён рядом с программой"
            initial_save_color = ft.Colors.GREY_600
        elif self.config.get("save_directory", ""):
            self.selected_save_directory = self.config.get("save_directory", "")
            initial_save_text = f"Папка сохранения: {self.selected_save_directory}"
            initial_save_color = ft.Colors.GREEN_700
        else:
            initial_save_text = "Выберите папку для сохранения"
            initial_save_color = ft.Colors.ORANGE_700

        self.save_directory_text = ft.Text(
            value=initial_save_text,
            color=initial_save_color
        )

        self.generate_btn = ft.ElevatedButton(
            "Создать DOCX документ",
            icon=ft.Icons.DESCRIPTION,
            on_click=self.generate_document,
            disabled=True,
            bgcolor=ft.Colors.GREY_400,
            color=ft.Colors.GREY_700,
            width=300,
            height=50,
            opacity=0.6,
            animate_opacity=300
        )

        def open_repo_on_footer_click(_e):
            self.page.launch_url(self.repo_url)

        footer = ft.Container(
            content=ft.Row([
                ft.CircleAvatar(
                    foreground_image_src=self.avatar_url,
                    radius=12,
                    bgcolor=ft.Colors.BLUE_700
                ),
                ft.Text(
                    "Made with ❤️ by Vennilay",
                    size=12,
                    color=ft.Colors.GREY_600,
                    italic=True
                )
            ], alignment=ft.MainAxisAlignment.CENTER, spacing=8),
            padding=ft.padding.only(top=20, bottom=10),
            on_click=open_repo_on_footer_click,
            tooltip="Открыть репозиторий на GitHub"
        )

        main_column = ft.Column([
            header_row,
            subtitle,
            ft.Divider(height=20, color=ft.Colors.BLUE_200),

            ft.Text("Данные титульного листа:", size=16, weight=ft.FontWeight.BOLD),
            self.group_field,
            self.student_field,
            self.teacher_field,
            self.work_number_field,

            ft.Divider(height=20, color=ft.Colors.BLUE_200),

            ft.Text("Настройки шаблона:", size=16, weight=ft.FontWeight.BOLD),
            self.template_path_field,
            ft.Row([select_template_btn, download_template_btn], spacing=10),
            self.template_path_display,

            ft.Divider(height=20, color=ft.Colors.BLUE_200),

            ft.Text("Дата документа:", size=16, weight=ft.FontWeight.BOLD),
            ft.Row([date_picker_btn, self.date_display], spacing=20),

            ft.Divider(height=20, color=ft.Colors.BLUE_200),

            ft.Text("Выбор файлов с кодом:", size=16, weight=ft.FontWeight.BOLD),
            select_dir_btn,
            self.directory_text,
            self.files_count_text,
            self.show_files_btn,

            ft.Divider(height=20, color=ft.Colors.BLUE_200),

            ft.Text("Место сохранения документа:", size=16, weight=ft.FontWeight.BOLD),
            ft.Row([self.save_nearby_checkbox], spacing=10),
            self.select_save_dir_btn,
            self.save_directory_text,

            ft.Divider(height=20, color=ft.Colors.BLUE_200),

            self.generate_btn,

            footer
        ], spacing=10)

        self.page.add(
            ft.Container(
                content=main_column,
                padding=20
            )
        )

        self.validate_form()

    def open_date_picker(self, _e):
        """Открывает календарь для выбора даты"""
        self.page.open(self.date_picker)

    @staticmethod
    def format_date(date: datetime) -> str:
        """Форматирует дату в нужный формат: «13» ноября 2025"""
        months = {
            1: "января", 2: "февраля", 3: "марта", 4: "апреля",
            5: "мая", 6: "июня", 7: "июля", 8: "августа",
            9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
        }
        return f"«{date.day}» {months[date.month]} {date.year}"

    def on_date_changed(self, event):
        """Обработчик изменения даты в календаре"""
        if event.control.value:
            self.selected_date = event.control.value
            self.date_display.value = self.format_date(self.selected_date)
            self.page.update()
            self.show_snackbar(
                f"✅ Дата выбрана: {self.format_date(self.selected_date)}",
                ft.Colors.GREEN_700
            )

    def on_date_dismissed(self, _e):
        """Обработчик закрытия календаря"""
        pass

    def find_code_files(self):
        """Ищет файлы с кодом в выбранной директории"""
        if not self.selected_directory:
            return

        extensions = ['.py', '.cpp', '.c', '.h', '.hpp', '.java', '.js', '.kt', '.go', '.rs']
        self.found_files = []

        try:
            for root, dirs, files in os.walk(self.selected_directory):
                for file in files:
                    if any(file.endswith(ext) for ext in extensions):
                        full_path = os.path.join(root, file)
                        self.found_files.append(full_path)
        except Exception as e:
            self.show_snackbar(
                f"❌ Ошибка при поиске файлов: {str(e)}",
                ft.Colors.RED_700
            )
            return

        if self.found_files:
            self.files_count_text.value = f"📁 Найдено файлов: {len(self.found_files)}"
            self.files_count_text.color = ft.Colors.GREEN_700
            self.files_count_text.weight = ft.FontWeight.BOLD
            self.show_files_btn.visible = True
        else:
            self.files_count_text.value = "❌ Файлы с кодом не найдены"
            self.files_count_text.color = ft.Colors.ORANGE_700
            self.show_files_btn.visible = False

        self.validate_form()

    def generate_document(self, _e):
        """Генерирует DOCX документ с титульным листом и кодом"""
        try:
            if not self.group_field.value:
                self.show_dialog("Ошибка", "Заполните поле 'Группа'!")
                return

            if not self.student_field.value:
                self.show_dialog("Ошибка", "Заполните поле 'ФИО студента'!")
                return

            if not self.teacher_field.value:
                self.show_dialog("Ошибка", "Заполните поле 'ФИО преподавателя'!")
                return

            if not self.work_number_field.value:
                self.show_dialog("Ошибка", "Заполните поле 'Номер работы'!")
                return

            if not self.found_files:
                self.show_dialog(
                    "Ошибка",
                    "Не выбраны файлы с кодом! Выберите директорию с файлами."
                )
                return

            template_path = self.template_path_field.value.strip()
            if not template_path:
                template_path = "template.docx"

            if not os.path.exists(template_path):
                self.show_dialog(
                    "Ошибка",
                    f"Файл шаблона не найден: {template_path}\n\n"
                    f"Убедитесь, что файл существует или укажите правильный путь."
                )
                return

            self.show_snackbar("⏳ Создание документа...", ft.Colors.BLUE_700)

            doc = DocxTemplate(template_path)
            context = {
                'group': self.group_field.value,
                'student_name': self.student_field.value,
                'teacher_name': self.teacher_field.value,
                'work_number': self.work_number_field.value,
                'date': self.format_date(self.selected_date)
            }
            doc.render(context)

            temp_file = "temp_output.docx"
            doc.save(temp_file)

            final_doc = Document(temp_file)

            for idx, file_path in enumerate(self.found_files, 1):
                if idx > 1:
                    final_doc.add_page_break()

                heading = final_doc.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                heading.paragraph_format.first_line_indent = Cm(1.25)
                heading.paragraph_format.space_before = Pt(0)
                heading.paragraph_format.space_after = Pt(6)

                run = heading.add_run(f"Задание № {idx}:")
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = True

                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        code_content = f.read()
                except Exception as read_error:
                    code_content = (
                        f"[Ошибка чтения файла: {os.path.basename(file_path)}\n"
                        f"Причина: {str(read_error)}]"
                    )

                code_para = final_doc.add_paragraph()
                code_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                code_para.paragraph_format.left_indent = Cm(1.25)
                code_para.paragraph_format.first_line_indent = Cm(0)
                code_para.paragraph_format.space_before = Pt(0)
                code_para.paragraph_format.space_after = Pt(0)
                code_para.paragraph_format.line_spacing = 1.0

                code_run = code_para.add_run(code_content)
                code_run.font.name = "Courier New"
                code_run.font.size = Pt(9)

            output_filename = (
                f"Отчёт_по_практической_работе_№{self.work_number_field.value}_"
                f"{self.student_field.value.replace(' ', '_')}.docx"
            )

            if self.save_nearby_checkbox.value:
                output_path = output_filename
            else:
                if not self.selected_save_directory:
                    self.show_dialog(
                        "Ошибка",
                        "Не выбрана папка для сохранения!\n\n"
                        "Выберите папку или включите опцию 'Сохранить рядом с программой'."
                    )
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                    return
                output_path = os.path.join(self.selected_save_directory, output_filename)

            final_doc.save(output_path)

            if os.path.exists(temp_file):
                os.remove(temp_file)

            self.save_config()

            absolute_path = os.path.abspath(output_path)
            self.show_dialog(
                "Успех! 🎉",
                f"Документ успешно создан!\n\n"
                f"Имя файла: {os.path.basename(output_path)}\n\n"
                f"Путь: {absolute_path}"
            )
            self.show_snackbar(
                f"✅ Документ создан: {os.path.basename(output_path)}",
                ft.Colors.GREEN_700
            )

        except Exception as ex:
            error_message = f"Произошла ошибка при создании документа:\n\n{str(ex)}"
            self.show_dialog("Ошибка", error_message)
            self.show_snackbar(f"❌ Ошибка: {str(ex)}", ft.Colors.RED_700)


def main(page: ft.Page):
    MireaReportGenerator(page)


if __name__ == "__main__":
    ft.app(target=main)
