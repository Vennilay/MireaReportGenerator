"""
Модуль для генерации DOCX документов
"""

import os
import tempfile
from datetime import datetime
from typing import List
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from utils.date_utils import format_date_russian
from utils.logger import logger


class DocumentGenerator:
    """
    Генератор DOCX документов
    """

    def __init__(self, template_path: str):
        self.template_path = template_path
        logger.log_operation("Инициализация DocumentGenerator", f"Шаблон: {template_path}")

    def validate_template(self) -> bool:
        exists = os.path.exists(self.template_path)
        if exists:
            logger.debug(f"Шаблон найден: {self.template_path}")
        else:
            logger.warning(f"Шаблон не найден: {self.template_path}")
        return exists

    def generate(
            self,
            group: str,
            student_name: str,
            teacher_name: str,
            work_number: str,
            date: datetime,
            code_files: List[str],
            output_path: str,
    ) -> bool:
        logger.info("=" * 70)
        logger.log_operation("Начало генерации документа")
        logger.info(f"  Группа: {group}")
        logger.info(f"  Студент: {student_name}")
        logger.info(f"  Преподаватель: {teacher_name}")
        logger.info(f"  Номер работы: {work_number}")
        logger.info(f"  Дата: {format_date_russian(date)}")
        logger.info(f"  Количество файлов кода: {len(code_files)}")
        logger.info(f"  Путь сохранения: {output_path}")

        try:
            logger.info("Загрузка шаблона DOCX...")
            doc = DocxTemplate(self.template_path)
            logger.log_file_operation("Загрузка шаблона", self.template_path, "успешно")

            context = {
                "group": group,
                "student_name": student_name,
                "teacher_name": teacher_name,
                "work_number": work_number,
                "date": format_date_russian(date),
            }
            logger.debug(f"Контекст для рендеринга: {context}")

            logger.info("Рендеринг титульного листа...")
            doc.render(context)
            logger.log_operation("Титульный лист отрендерен")

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                temp_file = tmp.name
            doc.save(temp_file)
            logger.debug(f"Сохранение временного файла: {temp_file}")
            logger.log_file_operation("Создание временного файла", temp_file, "успешно")

            logger.info("Открытие документа для добавления файлов кода...")
            final_doc = Document(temp_file)

            for idx, file_path in enumerate(code_files, 1):
                logger.info(f"Обработка файла {idx}/{len(code_files)}: {os.path.basename(file_path)}")

                if idx > 1:
                    logger.debug(f"  └─ Добавление разрыва страницы перед заданием {idx}")
                    final_doc.add_page_break()

                self._add_task_heading(final_doc, idx)
                logger.debug(f"  └─ Заголовок 'Задание № {idx}' добавлен")

                code_content = self._read_code_file(file_path)
                if "[Ошибка чтения файла" in code_content:
                    logger.warning(f"  └─ Ошибка чтения: {file_path}")
                else:
                    code_lines = len(code_content.splitlines())
                    code_size = len(code_content)
                    logger.debug(f"  └─ Прочитано {code_lines} строк, {code_size} байт")

                self._add_code_content(final_doc, code_content)
                logger.log_file_operation("Добавление кода", file_path, "успешно")

            logger.info(f"Сохранение финального документа: {output_path}")
            final_doc.save(output_path)
            logger.log_file_operation("Сохранение документа", output_path, "успешно")

            if os.path.exists(temp_file):
                logger.debug(f"Удаление временного файла: {temp_file}")
                os.remove(temp_file)
                logger.log_file_operation("Удаление временного файла", temp_file, "успешно")

            logger.info("=" * 70)
            logger.log_operation("Генерация документа завершена успешно", f"Файл: {output_path}")
            logger.info("=" * 70)
            return True

        except Exception as e:
            logger.error("=" * 70)
            logger.log_exception("Генерация документа", e)
            logger.error("=" * 70)
            print(f"Ошибка генерации документа: {str(e)}")
            return False

    @staticmethod
    def _add_task_heading(doc, task_number: int):
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        heading.paragraph_format.first_line_indent = Cm(1.25)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)

        run = heading.add_run(f"Задание № {task_number}:")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True

    @staticmethod
    def _read_code_file(file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                return content
        except Exception as e:
            error_msg = (
                f"[Ошибка чтения файла: {os.path.basename(file_path)}\n"
                f"Причина: {str(e)}]"
            )
            logger.error(f"Не удалось прочитать файл {file_path}: {str(e)}")
            return error_msg

    @staticmethod
    def _add_code_content(doc, code_content: str):
        code_para = doc.add_paragraph()
        code_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        code_para.paragraph_format.left_indent = Cm(1.25)
        code_para.paragraph_format.first_line_indent = Cm(0)
        code_para.paragraph_format.space_before = Pt(0)
        code_para.paragraph_format.space_after = Pt(0)
        code_para.paragraph_format.line_spacing = 1.0

        code_run = code_para.add_run(code_content)
        code_run.font.name = "Courier New"
        code_run.font.size = Pt(9)

    @staticmethod
    def generate_filename(work_number: str, student_name: str) -> str:
        safe_name = student_name.replace(" ", "_")
        filename = f"Отчёт_по_практической_работе_№{work_number}_{safe_name}.docx"
        logger.debug(f"Сгенерировано имя файла: {filename}")
        return filename