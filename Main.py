import os
import json
from datetime import datetime
import flet as ft
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class MireaReportGenerator:
    def __init__(self, page: ft.Page):
        self.page = page
        self.page.title = "MIREA Report Generator - Генератор отчётов РТУ МИРЭА"
        self.page.window.width = 850
        self.page.window.height = 1000
        self.page.padding = 20

        self.page.locale_configuration = ft.LocaleConfiguration(
            supported_locales=[ft.Locale("ru", "RU")],
            current_locale=ft.Locale("ru", "RU")
        )

        self.selected_directory = None
        self.found_files = []
        self.config_file = "config.json"
        self.selected_date = datetime.now()

        self.group_field = None
        self.student_field = None
        self.teacher_field = None
        self.work_number_field = None
        self.template_path_field = None
        self.date_picker = None
        self.date_display = None
        self.directory_text = None
        self.template_path_display = None
        self.file_picker = None
        self.template_file_picker = None
        self.files_list = None

        self.config = self.load_config()
        self.create_ui()

    def load_config(self):
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            self.show_snackbar(f"Ошибка загрузки конфига: {str(e)}", ft.Colors.ORANGE)

        return {
            "group": "",
            "student_name": "",
            "teacher_name": "",
            "work_number": "",
            "last_directory": "",
            "template_path": "template.docx"
        }

    def save_config(self):
        try:
            config = {
                "group": self.group_field.value,
                "student_name": self.student_field.value,
                "teacher_name": self.teacher_field.value,
                "work_number": self.work_number_field.value,
                "last_directory": self.selected_directory or "",
                "template_path": self.template_path_field.value
            }
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config, indent=4, fp=f, ensure_ascii=False)
        except Exception as e:
            self.show_snackbar(f"Ошибка сохранения конфига: {str(e)}", ft.Colors.ORANGE)

    def show_snackbar(self, message: str, color: str = ft.Colors.BLUE_700):
        snackbar = ft.SnackBar(
            content=ft.Text(message, color=ft.Colors.WHITE),
            bgcolor=color,
            duration=3000
        )
        self.page.overlay.append(snackbar)
        snackbar.open = True
        self.page.update()

    def show_dialog(self, title: str, message: str):
        def close_dialog(_e):
            dialog.open = False
            self.page.update()

        dialog = ft.AlertDialog(
            modal=True,
            title=ft.Text(title, weight=ft.FontWeight.BOLD),
            content=ft.Text(message),
            actions=[ft.TextButton("ОК", on_click=close_dialog)],
            actions_alignment=ft.MainAxisAlignment.END
        )

        self.page.overlay.append(dialog)
        dialog.open = True
        self.page.update()

    def show_about_dialog(self, _e):
        def close_dialog(_e):
            dialog.open = False
            self.page.update()

        def open_github(_e):
            self.page.launch_url("https://github.com/Vennilay")

        dialog = ft.AlertDialog(
            modal=True,
            title=ft.Text("О создателе 👨‍💻", weight=ft.FontWeight.BOLD, size=20),
            content=ft.Column([
                ft.Text("MIREA Report Generator", size=16, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_700),
                ft.Divider(height=10),
                ft.Text("Генератор отчётов для студентов РТУ МИРЭА", size=14, color=ft.Colors.GREY_700),
                ft.Container(height=10),
                ft.Text("Разработчик: Vennilay", size=14, weight=ft.FontWeight.W_500),
                ft.Container(height=5),
                ft.ElevatedButton(
                    "GitHub Profile", icon=ft.Icons.OPEN_IN_NEW, on_click=open_github,
                    style=ft.ButtonStyle(bgcolor=ft.Colors.GREY_800, color=ft.Colors.WHITE)
                ),
                ft.Container(height=10),
                ft.Text("© 2025 Vennilay", size=12, color=ft.Colors.GREY_500, italic=True)
            ], tight=True, spacing=5),
            actions=[ft.TextButton("Закрыть", on_click=close_dialog)],
            actions_alignment=ft.MainAxisAlignment.END
        )

        self.page.overlay.append(dialog)
        dialog.open = True
        self.page.update()

    def create_ui(self):
        header_row = ft.Row([
            ft.Text("MIREA Report Generator", size=26, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_700),
            ft.IconButton(icon=ft.Icons.INFO_OUTLINED, tooltip="О создателе",
                          on_click=self.show_about_dialog, icon_color=ft.Colors.BLUE_600, icon_size=28)
        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN)

        subtitle = ft.Text("Генератор отчётов для РТУ МИРЭА", size=14, color=ft.Colors.GREY_600, italic=True)

        self.group_field = ft.TextField(
            label="Группа (например: ИКБО-47-52)",
            value=self.config.get("group", ""), width=400,
            autofocus=True, border_color=ft.Colors.BLUE_400,
            prefix_icon=ft.Icons.GROUP
        )

        self.student_field = ft.TextField(
            label="ФИО студента (например: Иванов И.И.)",
            value=self.config.get("student_name", ""), width=400,
            border_color=ft.Colors.BLUE_400, prefix_icon=ft.Icons.PERSON
        )

        self.teacher_field = ft.TextField(
            label="ФИО преподавателя", value=self.config.get("teacher_name", ""),
            width=400, border_color=ft.Colors.BLUE_400, prefix_icon=ft.Icons.SCHOOL
        )

        self.work_number_field = ft.TextField(
            label="Номер работы", value=self.config.get("work_number", ""),
            width=200, keyboard_type=ft.KeyboardType.NUMBER,
            border_color=ft.Colors.BLUE_400, prefix_icon=ft.Icons.NUMBERS
        )

        self.template_path_field = ft.TextField(
            label="Путь к файлу шаблона",
            value=self.config.get("template_path", "template.docx"),
            width=400, border_color=ft.Colors.BLUE_400,
            prefix_icon=ft.Icons.DESCRIPTION,
            hint_text="Укажите путь или имя файла шаблона"
        )

        self.template_file_picker = ft.FilePicker(on_result=self.on_template_selected)
        self.page.overlay.append(self.template_file_picker)

        select_template_btn = ft.ElevatedButton(
            "Выбрать файл шаблона", icon=ft.Icons.FILE_OPEN,
            on_click=lambda _: self.template_file_picker.pick_files(
                dialog_title="Выберите файл шаблона DOCX",
                allowed_extensions=["docx"], file_type=ft.FilePickerFileType.CUSTOM
            ),
            style=ft.ButtonStyle(bgcolor=ft.Colors.PURPLE_600, color=ft.Colors.WHITE)
        )

        self.template_path_display = ft.Text(
            value=f"Текущий шаблон: {self.config.get('template_path', 'template.docx')}",
            color=ft.Colors.GREY_700, size=12
        )

        self.date_picker = ft.DatePicker(
            first_date=datetime(2020, 1, 1),
            last_date=datetime(2030, 12, 31),
            on_change=self.on_date_changed,
            on_dismiss=self.on_date_dismissed,
            help_text="Выберите дату",
            cancel_text="Отмена",
            confirm_text="ОК",
            error_format_text="Неверный формат",
            error_invalid_text="Вне диапазона",
            field_label_text="Введите дату",
            field_hint_text="дд.мм.гггг"
        )
        self.page.overlay.append(self.date_picker)

        self.date_display = ft.Text(
            value=self.format_date(self.selected_date),
            size=16, color=ft.Colors.GREEN_700, weight=ft.FontWeight.BOLD
        )

        date_picker_btn = ft.ElevatedButton(
            "Выбрать дату", icon=ft.Icons.CALENDAR_MONTH,
            on_click=self.open_date_picker,
            style=ft.ButtonStyle(bgcolor=ft.Colors.BLUE_600, color=ft.Colors.WHITE)
        )

        self.directory_text = ft.Text(value="Директория не выбрана", color=ft.Colors.GREY_700)

        self.file_picker = ft.FilePicker(on_result=self.on_directory_selected)
        self.page.overlay.append(self.file_picker)

        select_dir_btn = ft.ElevatedButton(
            "Выбрать директорию с кодом", icon=ft.Icons.FOLDER_OPEN,
            on_click=lambda _: self.file_picker.get_directory_path(
                dialog_title="Выберите папку с файлами кода"
            ),
            style=ft.ButtonStyle(bgcolor=ft.Colors.BLUE_600, color=ft.Colors.WHITE)
        )

        self.files_list = ft.Column(spacing=5)

        generate_btn = ft.ElevatedButton(
            "Создать DOCX документ",
            icon=ft.Icons.DESCRIPTION,
            on_click=self.generate_document,
            style=ft.ButtonStyle(bgcolor=ft.Colors.GREEN_700, color=ft.Colors.WHITE),
            width=300, height=50
        )

        footer = ft.Container(
            content=ft.Row([
                ft.Text("Made with ❤️ by Vennilay", size=12,
                        color=ft.Colors.GREY_600, italic=True)
            ], alignment=ft.MainAxisAlignment.CENTER),
            padding=ft.padding.only(top=20, bottom=10)
        )

        self.page.add(
            ft.Container(
                content=ft.Column([
                    header_row,
                    subtitle,
                    ft.Divider(height=20, color=ft.Colors.BLUE_200),

                    ft.Text("Данные титульного листа:", size=16, weight=ft.FontWeight.BOLD),
                    self.group_field,
                    self.student_field,
                    self.teacher_field,
                    self.work_number_field,

                    ft.Divider(height=20, color=ft.Colors.BLUE_200),

                    ft.Text("Настройки шаблона:", size=16, weight=ft.FontWeight.BOLD),
                    ft.Row([self.template_path_field, select_template_btn], spacing=10),
                    self.template_path_display,

                    ft.Divider(height=20, color=ft.Colors.BLUE_200),

                    ft.Text("Дата документа:", size=16, weight=ft.FontWeight.BOLD),
                    ft.Row([date_picker_btn, self.date_display], spacing=20),

                    ft.Divider(height=20, color=ft.Colors.BLUE_200),

                    ft.Text("Выбор файлов с кодом:", size=16, weight=ft.FontWeight.BOLD),
                    select_dir_btn,
                    self.directory_text,
                    self.files_list,

                    ft.Divider(height=20, color=ft.Colors.BLUE_200),

                    generate_btn,

                    footer
                ], spacing=10, scroll=ft.ScrollMode.AUTO),
                padding=20
            )
        )

    def on_template_selected(self, e: ft.FilePickerResultEvent):
        if e.files and len(e.files) > 0:
            selected_file = e.files[0]
            template_path = selected_file.path

            self.template_path_field.value = template_path
            self.template_path_display.value = f"Текущий шаблон: {os.path.basename(template_path)}"
            self.template_path_display.color = ft.Colors.GREEN_700

            self.page.update()
            self.show_snackbar(f"✅ Выбран шаблон: {os.path.basename(template_path)}", ft.Colors.GREEN_700)

    def open_date_picker(self, _e):
        self.page.open(self.date_picker)

    @staticmethod
    def format_date(date: datetime) -> str:
        months = {
            1: "января", 2: "февраля", 3: "марта", 4: "апреля",
            5: "мая", 6: "июня", 7: "июля", 8: "августа",
            9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
        }
        return f"«{date.day}» {months[date.month]} {date.year}"

    def on_date_changed(self, event):
        if event.control.value:
            self.selected_date = event.control.value
            self.date_display.value = self.format_date(self.selected_date)
            self.page.update()
            self.show_snackbar(f"✅ Дата выбрана: {self.format_date(self.selected_date)}",
                               ft.Colors.GREEN_700)

    def on_date_dismissed(self, _e):
        pass

    def on_directory_selected(self, e: ft.FilePickerResultEvent):
        if e.path:
            self.selected_directory = e.path
            self.directory_text.value = f"Выбрана: {self.selected_directory}"
            self.directory_text.color = ft.Colors.GREEN_700

            self.find_code_files()
            self.page.update()
            self.show_snackbar(f"✅ Найдено файлов: {len(self.found_files)}",
                               ft.Colors.GREEN_700)

    def find_code_files(self):
        if not self.selected_directory:
            return

        extensions = ['.py', '.cpp', '.c', '.h', '.hpp', '.java', '.js', '.kt', '.go', '.rs']
        self.found_files = []

        try:
            for root, dirs, files in os.walk(self.selected_directory):
                for file in files:
                    if any(file.endswith(ext) for ext in extensions):
                        full_path = os.path.join(root, file)
                        self.found_files.append(full_path)
        except Exception as e:
            self.show_snackbar(f"❌ Ошибка при поиске файлов: {str(e)}", ft.Colors.RED_700)
            return

        self.files_list.controls.clear()

        if self.found_files:
            self.files_list.controls.append(
                ft.Text(f"Найдено файлов: {len(self.found_files)}",
                        weight=ft.FontWeight.BOLD,
                        color=ft.Colors.GREEN_700)
            )
            for file_path in self.found_files:
                self.files_list.controls.append(
                    ft.Text(f"  • {os.path.basename(file_path)}", color=ft.Colors.GREY_700)
                )
        else:
            self.files_list.controls.append(
                ft.Text("Файлы с кодом не найдены", color=ft.Colors.RED_700)
            )

    def generate_document(self, _e):
        try:
            if not self.group_field.value:
                self.show_dialog("Ошибка", "Заполните поле 'Группа'!")
                return

            if not self.student_field.value:
                self.show_dialog("Ошибка", "Заполните поле 'ФИО студента'!")
                return

            if not self.teacher_field.value:
                self.show_dialog("Ошибка", "Заполните поле 'ФИО преподавателя'!")
                return

            if not self.work_number_field.value:
                self.show_dialog("Ошибка", "Заполните поле 'Номер работы'!")
                return

            if not self.found_files:
                self.show_dialog("Ошибка", "Не выбраны файлы с кодом! Выберите директорию с файлами.")
                return

            template_path = self.template_path_field.value.strip()
            if not template_path:
                template_path = "template.docx"

            if not os.path.exists(template_path):
                self.show_dialog(
                    "Ошибка",
                    f"Файл шаблона не найден: {template_path}\n\n"
                    f"Убедитесь, что файл существует или укажите правильный путь."
                )
                return

            self.show_snackbar("⏳ Создание документа...", ft.Colors.BLUE_700)

            doc = DocxTemplate(template_path)
            context = {
                'group': self.group_field.value,
                'student_name': self.student_field.value,
                'teacher_name': self.teacher_field.value,
                'work_number': self.work_number_field.value,
                'date': self.format_date(self.selected_date)
            }
            doc.render(context)

            temp_file = "temp_output.docx"
            doc.save(temp_file)

            final_doc = Document(temp_file)

            for idx, file_path in enumerate(self.found_files, 1):
                if idx > 1:
                    final_doc.add_page_break()

                heading = final_doc.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                heading.paragraph_format.first_line_indent = Cm(1.25)
                heading.paragraph_format.space_before = Pt(0)
                heading.paragraph_format.space_after = Pt(6)

                run = heading.add_run(f"Задание № {idx}:")
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = True

                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        code_content = f.read()
                except Exception as read_error:
                    code_content = (
                        f"[Ошибка чтения файла: {os.path.basename(file_path)}\n"
                        f"Причина: {str(read_error)}]"
                    )

                code_para = final_doc.add_paragraph()
                code_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                code_para.paragraph_format.left_indent = Cm(1.25)
                code_para.paragraph_format.first_line_indent = Cm(0)
                code_para.paragraph_format.space_before = Pt(0)
                code_para.paragraph_format.space_after = Pt(0)
                code_para.paragraph_format.line_spacing = 1.0

                code_run = code_para.add_run(code_content)
                code_run.font.name = "Courier New"
                code_run.font.size = Pt(9)

            output_filename = (
                f"Работа_{self.work_number_field.value}_"
                f"{self.student_field.value.replace(' ', '_')}.docx"
            )
            final_doc.save(output_filename)

            if os.path.exists(temp_file):
                os.remove(temp_file)

            self.save_config()

            absolute_path = os.path.abspath(output_filename)
            self.show_dialog(
                "Успех! 🎉",
                f"Документ успешно создан!\n\n"
                f"Имя файла: {output_filename}\n\n"
                f"Путь: {absolute_path}"
            )
            self.show_snackbar(f"✅ Документ создан: {output_filename}", ft.Colors.GREEN_700)

        except Exception as ex:
            error_message = f"Произошла ошибка при создании документа:\n\n{str(ex)}"
            self.show_dialog("Ошибка", error_message)
            self.show_snackbar(f"❌ Ошибка: {str(ex)}", ft.Colors.RED_700)


def main(page: ft.Page):
    MireaReportGenerator(page)


if __name__ == "__main__":
    ft.app(target=main)